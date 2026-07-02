"""
File content extraction module.

Extracts text content from various file formats (PDF, DOCX, TXT, CSV).
"""

from typing import Optional, Dict, Any, List
from pathlib import Path
import csv
import io

from core.logger import get_logger

logger = get_logger(__name__)


class FileExtractor:
    """
    Extracts text content from downloaded files.
    
    Supports PDF, DOCX, TXT, and CSV formats.
    """

    def __init__(self) -> None:
        """Initialize the file extractor."""
        pass

    def extract(self, file_path: Path, file_type: str) -> Optional[str]:
        """
        Extract content from file.
        
        Args:
            file_path: Path to file
            file_type: File type (pdf, docx, txt, csv)
            
        Returns:
            Extracted text content or None
        """
        try:
            file_type = file_type.lower()
            
            if file_type == "pdf":
                return self._extract_pdf(file_path)
            elif file_type in ("docx", "doc"):
                return self._extract_docx(file_path)
            elif file_type == "txt":
                return self._extract_txt(file_path)
            elif file_type == "csv":
                return self._extract_csv(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None
        
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return None

    def _extract_pdf(self, file_path: Path) -> Optional[str]:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            import PyPDF2
            
            text_parts = []
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                logger.debug(f"Extracting PDF with {num_pages} pages")
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text:
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}\n")
            
            return "\n".join(text_parts) if text_parts else None
        
        except ImportError:
            logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
            return None
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            return None

    def _extract_docx(self, file_path: Path) -> Optional[str]:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract table content
            for table in doc.tables:
                text_parts.append("\n| Table |\n|-------|")
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    text_parts.append("| " + " | ".join(cells) + " |")
            
            return "\n\n".join(text_parts) if text_parts else None
        
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return None

    def _extract_txt(self, file_path: Path) -> Optional[str]:
        """
        Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            if content.strip():
                return content
            return None
        
        except Exception as e:
            logger.error(f"Error extracting TXT: {e}")
            return None

    def _extract_csv(self, file_path: Path) -> Optional[str]:
        """
        Extract content from CSV file.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            CSV content as markdown table
        """
        try:
            rows = []
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                csv_reader = csv.reader(f)
                rows = list(csv_reader)
            
            if not rows:
                return None
            
            # Convert to markdown table
            markdown_parts = []
            
            # Header row
            if len(rows) > 0:
                header = rows[0]
                markdown_parts.append("| " + " | ".join(header) + " |")
                markdown_parts.append("| " + " | ".join(["---"] * len(header)) + " |")
            
            # Data rows
            for row in rows[1:]:
                markdown_parts.append("| " + " | ".join(row) + " |")
            
            return "\n".join(markdown_parts)
        
        except Exception as e:
            logger.error(f"Error extracting CSV: {e}")
            return None

    def extract_metadata(self, file_path: Path, file_type: str) -> Dict[str, Any]:
        """
        Extract metadata from file.
        
        Args:
            file_path: Path to file
            file_type: File type
            
        Returns:
            Metadata dictionary
        """
        try:
            file_stat = file_path.stat()
            
            metadata = {
                "filename": file_path.name,
                "file_type": file_type.lower(),
                "file_size": file_stat.st_size,
                "modified_time": file_stat.st_mtime,
            }
            
            # Type-specific metadata
            if file_type.lower() == "pdf":
                metadata.update(self._get_pdf_metadata(file_path))
            elif file_type.lower() in ("docx", "doc"):
                metadata.update(self._get_docx_metadata(file_path))
            elif file_type.lower() == "csv":
                metadata.update(self._get_csv_metadata(file_path))
            
            return metadata
        
        except Exception as e:
            logger.warning(f"Error extracting metadata: {e}")
            return {"filename": file_path.name, "file_type": file_type}

    def _get_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF metadata."""
        try:
            import PyPDF2
            
            metadata = {}
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                if pdf_reader.metadata:
                    metadata["pages"] = len(pdf_reader.pages)
                    if pdf_reader.metadata.title:
                        metadata["title"] = pdf_reader.metadata.title
                    if pdf_reader.metadata.author:
                        metadata["author"] = pdf_reader.metadata.author
            
            return metadata
        except Exception as e:
            logger.debug(f"Error getting PDF metadata: {e}")
            return {}

    def _get_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX metadata."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            metadata = {}
            
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
            
            # Count paragraphs and tables
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            
            return metadata
        except Exception as e:
            logger.debug(f"Error getting DOCX metadata: {e}")
            return {}

    def _get_csv_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract CSV metadata."""
        try:
            metadata = {}
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                csv_reader = csv.reader(f)
                rows = list(csv_reader)
            
            if rows:
                metadata["row_count"] = len(rows)
                metadata["column_count"] = len(rows[0]) if rows else 0
                metadata["columns"] = rows[0] if rows else []
            
            return metadata
        except Exception as e:
            logger.debug(f"Error getting CSV metadata: {e}")
            return {}
